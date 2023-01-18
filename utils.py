import time
import os

from tkinter import filedialog

from docx import Document
from docx.shared import Inches


def seleciona_imagens(quant: str):
    document = Document()
    # janela_padrao = Tk().withdraw()
    try:
        aux_quant = int(quant)
        for i in range(aux_quant):
            num_foto = str(i + 1)
            if i + 1 < 10:
                num_foto = "0" + str(i + 1)
            document.add_paragraph(
                'Foto ' + num_foto + ":"
            )
            a = filedialog.askopenfile()
            document.add_picture(a.name, width=Inches(2.5))

    except Exception:
        print("Erro! Verifique se a quantidade de imagens foi informada.")

    name_file = "/arquivo_gerado_" + str(time.time()) + ".docx"
    caminho = "documentos_gerados"
    verificar_caminho_arquivo(caminho)
    document.save(caminho + name_file)

def verificar_caminho_arquivo(caminho):
    if not os.path.exists(caminho):
        os.makedirs(caminho)
